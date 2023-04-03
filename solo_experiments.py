import main_funks as m
import time
import docx
import matplotlib.pyplot as plt
from config import settings as s

# одиночные эксперименты в сложных местах 

if __name__ == '__main__':
    # 10, 11 элементов G_inh < -0.07

    # Необходимые параметры
    G_inh = -0.07
    s.k_systems = 15
    tMax = 500
    modifier = 'solo'

    # Пути
    path_x_start = s.Graphic_data_path + '_x' + modifier + '.png'
    path_x_end = s.Graphic_data_path + '_x' + modifier + '_end.png'
    path_R = s.Graphic_data_path + '_R' + modifier + '.png'
    path_graph_last_state = s.Graphic_data_path + '_lsFHN_' + modifier + '.png'

    m.generate_file_names_R12_Ginh(modifier=modifier)
    IC = m.generate_IC_any_sizes(dist_between_neurons=s.dist_between_neurons_IC, 
                                 type='full')
    m.plot_IC_FHN(IC, s.pathIC, text='Начальные условия')

    R1_arr, R2_arr, IC, depressed_elements, last_state, [xs, ys, zs, ts] = m.\
        make_experiment(G_inh, IC, tMax, s.highAccuracy, path_graph_x_start=path_x_start, path_graph_x_end=path_x_end, 
        path_graph_R=path_R, path_graph_last_state=path_graph_last_state, do_need_show=False, do_need_xyzt=True)


    mydoc = docx.Document()

    # Делаем первые НУ (по дефолту противофазные/циклопные)
    IC = m.generate_IC_any_sizes(dist_between_neurons=s.dist_between_neurons_IC, type=s.IC_type)
    m.plot_IC_FHN(IC, s.pathIC, text='Начальные условия')

    mydoc.add_heading("Solo experiment. G_inh = " + str(G_inh) + '; k_systems = ' + str(s.k_systems))
    mydoc.add_heading("Initial conditions:", 2)
    for j in range(0, s.k_systems):
        mydoc.add_paragraph(str(IC[j * s.k]) + ', ' + str(IC[j * s.k + 1]) + ', ' +
        str(IC[j * s.k + 2]) + ', ')
    mydoc.add_picture(s.pathIC)
    mydoc.add_page_break()

    mydoc.add_picture(path_x_start, width=docx.shared.Inches(6.5))
    mydoc.add_picture(path_x_end, width=docx.shared.Inches(6.5))
    mydoc.add_picture(path_R, width=docx.shared.Inches(5))

    mydoc.save(s.path_Doc)

    plt.figure(figsize=(18, 5))
    for i in range(s.k_systems):
        plt.plot(ts, xs[i], label=('eq' + str(i + 1)), 
                 linestyle=s.plot_styles[i], color=s.plot_colors[i])
        plt.legend()
    plt.xlabel('t')
    plt.ylabel('x')
    plt.savefig(s.Graphic_data_path + '_x' + '_full' + '.png')
    plt.grid()
    plt.show()
    plt.close()

    plt.figure(figsize=(5, 5))
    for i in range(s.k_systems):
        plt.plot(xs[i], ys[i], label=('eq' + str(i + 1)), 
                 linestyle=s.plot_styles[i], color=s.plot_colors[i])
        plt.legend()
    plt.xlabel('x')
    plt.ylabel('y')
    plt.savefig(s.Graphic_data_path + '_xy' + '_full' + '.png')
    plt.grid()
    plt.show()


    

