import main_funks as m
import memory_work as mem
import time
import docx
import joblib
import matplotlib.pyplot as plt
import numpy as np
from config import settings as s
from datetime import datetime


# Обертка для легкой работы с параллелем, основные изменения (параметр связи, пути сохранения) можно делать здесь
def ex_Ginh_f(index, IC, divider, do_need_show=False, type_taking_IC = 'rand', G_inh_ = '0', a = s.a):

    path_x_start = s.Graphic_data_path + '_x' + str(index) + '.png'
    path_x_end = s.Graphic_data_path + '_x' + str(index) + '_end.png'
    path_R = s.Graphic_data_path + '_R' + str(index) + '.png'
    path_graph_last_state = s.Graphic_data_path + '_lsFHN_' + str(index) + '.png'
    path_graph_phi_k = s.Graphic_data_path + '_phi_' + str(index) + '.png'
    path_IC = s.Graphic_data_path + '_IC_' + str(index) + '.png'
    
    if isinstance(IC, int or bool):
        if type_taking_IC == 'rand':
            IC = mem.IC_random_generator(-3, 3, pathSave=path_IC)
        elif type_taking_IC == 'FHN':
            IC = mem.IC_FHN_random_generator(path=s.FHN_tr_path, pathSave=path_IC)

    if G_inh_ == '0':
        G_inh = s.G_inh_sign * index / float(divider)
    else:
        G_inh = G_inh_

    if G_inh_ == '0':
        # Только для протяжки берем так
        if s.G_inh_sign < 0:
            if G_inh >= s.G_inh_sign * 0.001:
                tMax = s.tMax1
            elif G_inh >= s.G_inh_sign * 0.005:
                tMax = s.tMax2
            else:
                tMax = s.tMax3
        else:
            if s.G_inh_sign == 0:
                tMax = 1000
            if G_inh <= s.G_inh_sign * 0.001:
                tMax = s.tMax1
            elif G_inh <= s.G_inh_sign * 0.005:
                tMax = s.tMax2
            else:
                tMax = s.tMax3
    else:
        tMax = s.tMax3

    print('Experiment ' + str(index), 'tMax ' + str(tMax))
    R1_arr, R2_arr, IC, other = m.\
        make_experiment(G_inh, IC, tMax, s.highAccuracy, path_graph_x_start=path_x_start, path_graph_x_end=path_x_end,
                        path_graph_R=path_R, path_graph_last_state=path_graph_last_state, 
                        path_graph_phi=path_graph_phi_k, do_need_show=do_need_show, a=a)

    #other = depressed_elements, last_state, phi_arr

    # Возращает [*R, IC, *paths, G_inh, *other]
    return [R1_arr, R2_arr], IC, [path_x_start, path_x_end, path_R, path_graph_last_state, path_graph_phi_k, path_IC], G_inh, other, 

def make_protyazhka_R12_dep_G_inh_full(start, stop, step, divider=1000, small_modifier=''):
    # Настраиваем пути 
    s.G_inh_sign = start // abs(start)
    modifier = 'pr_full'
    if step > 0:
        small_modifier_ = 'left'
    else:
        small_modifier_ = 'right'
    small_modifier = small_modifier_ + small_modifier
    mem.generate_file_names_R12_Ginh(modifier, small_modifier=small_modifier)
    print('files: ', str(s.path_Doc), str(s.R12_last_path))
    r12_last_path = s.R12_last_path

    # Инициализируем
    final_time = time.time()
    # random_IC = mem.IC_random_generator(-3, 3)
    R1_arr_last = []
    R2_arr_last = []
    G_inh_arr = []
    norm_of_periods_arr = []

    # Делаем первые НУ (по дефолту противофазные/циклопные)
    IC = m.generate_IC_any_sizes(dist_between_neurons=s.dist_between_neurons_IC, type=s.IC_type, params_random_generation = (-3, 3, s.pathIC))
    mem.plot_IC_FHN(IC, s.pathIC, text='Начальные условия')

    # Добавляем неоднородность по параметру а 
    a_arr = m.make_heterogeneity_a(s.a, s.delta)


    # Инициализируем файл doc
    mydoc = docx.Document()
    # Записываем параметр а
    a_paragraph = mydoc.add_paragraph('a = [')
    for a in a_arr:
        a_paragraph.add_run(str(a) + ', ')
    a_paragraph.add_run(f']; delta = {s.delta}')
    mydoc.add_heading("Initial conditions:", 2)
    for j in range(0, s.k_systems):
        mydoc.add_paragraph(str(IC[j * s.k]) + ', ' + str(IC[j * s.k + 1]) + ', ' +
                            str(IC[j * s.k + 2]) + ', ')
    mydoc.add_picture(s.pathIC)
    mydoc.add_paragraph('Final graphs data in file: ' + str(s.R12_last_path))
    mydoc.add_page_break()

    for i in range(start, stop, step):
        if i >= 0:
            s.G_inh_sign = 1
        else:
            s.G_inh_sign = -1

        R, IC_i, paths, G_inh_i, other = ex_Ginh_f(abs(i), IC, divider, False, a = a_arr)

        R1_arr_i = R[0]
        R2_arr_i = R[1]
        path_x_start = paths[0]
        path_x_end = paths[1]
        path_R = paths[2]
        path_graph_last_state = paths[3]
        path_graph_phi_k = paths[4]
        depressed_elements_i = other[0]
        last_state_i = other[1]
        phi_arr_i = other[2]
        mean_periods, difference_of_mean_periods = other[3]
            
        # Запись в файл. Docx
        mydoc.add_heading('Experiment ' + str(i), 1)
        mydoc.add_picture(path_x_start, width=docx.shared.Inches(6.5))
        mydoc.add_picture(path_x_end, width=docx.shared.Inches(6.5))
        mydoc.add_picture(path_R, width=docx.shared.Inches(4))

        # Рисуем конечное состояние на единичной окружности
        mydoc.add_picture(path_graph_last_state, width=docx.shared.Inches(4))

        # Рисуем график phi(k)
        mydoc.add_picture(path_graph_phi_k, width=docx.shared.Inches(4.5))

        # Рисуем график изменения суммы фаз относительно первого элемента
        # phi_arr_i_T = phi_arr_i.T
        # sum_phi = []
        # for i in range(len(phi_arr_i[0])):
        #     sum_phi.append(sum(phi_arr_i_T[i]))

        # График сумм фаз
        # path_graph_sum_phi = s.Graphic_data_path + '_sumPhi_' + str(0) + '.png'
        # mem.draw_sum_phases(sum_phi, path_graph_sum_phi)
        # mydoc.add_picture(path_graph_sum_phi, width=docx.shared.Inches(6))

        # Записываем средние периоды
        mydoc.add_heading(f'max(mean_periods) - min(mean_periods) {difference_of_mean_periods}', 3)
        norm_of_periods_arr.append(difference_of_mean_periods)

        # Записываем подавленные элементы
        if len(depressed_elements_i) != 0:
            print('depressed', depressed_elements_i)
            mydoc.add_heading('Подавленные элементы: ', 3)
            for el in depressed_elements_i:
                mydoc.add_paragraph(str(el))

        mydoc.add_page_break()

        # Сохраняем в док результатов всё что сделали
        mydoc.save(s.path_Doc)

        # Последние элементы массива R1, R2
        R1_arr_last.append(R1_arr_i[-1])
        R2_arr_last.append(R2_arr_i[-1])
        G_inh_arr.append(G_inh_i)

        # Записываем
        mem.write_R12_last(G_inh_arr, R1_arr_last, R2_arr_last, r12_last_path)

        # Делаем начальное состояние следующего эксперемента как конечное состояние предыдущего
        IC = last_state_i
    
    # Делаем график зависимости параметров порядка от силы связи
    mem.draw_R_dep_G(G_inh_arr, R1_arr_last, R2_arr_last, path=s.graph_R_Ginh_path, modifier=modifier)

    # Добавляем этот график в док
    mydoc.add_heading('Final graph', 1)
    mydoc.add_picture(s.graph_R_Ginh_path, width=docx.shared.Inches(5.5))

    # График изменения нормы средних периодов от связи
    plt.plot(G_inh_arr, norm_of_periods_arr)
    plt.grid()
    plt.xlabel('G_inh')
    plt.ylabel('‖periods‖∞')
    norm_periods_fig_path = s.Graphic_data_path + '_norm_periods.png'
    plt.savefig(norm_periods_fig_path)
    plt.close()
    mydoc.add_picture(norm_periods_fig_path, width=docx.shared.Inches(5.5))


    mydoc.save(s.path_Doc)
    # Save data
    # mem.recordICAndR(R_data_path, R1_arr, IC_arr, G_inh_arr, index)
    print('Process end. Final time: ', time.time() - final_time)

def make_experiments_with_random_IC(G_inh_points_arr, n_exps_on_point = 50, type_taking_IC = 'FHN'):
    # Например для 8 элементов
    # G_inh_points_arr = [-0.062, -0.02, -0.01, 0.02, 0.06]

    # Инициализируем файл doc
    mydoc = docx.Document()

    # Генерируем пути для итоговых файлов с результатами
    modifier = 'statistic_rand_IC'
    small_modifier = str(n_exps_on_point)
    mem.generate_file_names_R12_Ginh(modifier, small_modifier=small_modifier)
    print('files: ', str(s.path_Doc), str(s.R12_last_path))
    r12_last_path = s.R12_last_path

    # Инициализируем
    final_time = time.time()

    n_existance_cycles = n_exps_on_point // s.n_streams

    print('Start experements time: ', datetime.now().time())

    for cycle_num, G_inh_k_exp in enumerate(G_inh_points_arr):
    
        # Инициализируем
        R1_arr_last = []
        R2_arr_last = []
        G_inh_arr = []
        
        mydoc.add_heading('Cycle num ' + str(cycle_num) + ', G_inh = ' + str(G_inh_k_exp), 1)
        for i in range(n_existance_cycles + 1):
            
            existance = 0
            n_streams = s.n_streams
            if i != n_existance_cycles:
                
                existance = joblib.Parallel(n_jobs=n_streams)(joblib.delayed(ex_Ginh_f)(k, 0, False, G_inh_=G_inh_k_exp)
                                                                    for k in range(i*n_streams, (i+1)*n_streams))
            elif n_exps_on_point%n_streams != 0:
                existance = joblib.Parallel(n_jobs=n_exps_on_point%n_streams)(joblib.delayed(ex_Ginh_f)(k, 0, False, G_inh_=G_inh_k_exp)
                                                                    for k in range(i*n_streams, i*n_streams+n_exps_on_point%n_streams, G_inh_=G_inh_k_exp))

            elif i == n_existance_cycles and n_exps_on_point%n_streams == 0:
                continue
            for ex_index, existance_j in enumerate(existance):
                R, IC_i, paths, G_inh_i, other = existance_j

                R1_arr_i = R[0]
                R2_arr_i = R[1]
                path_x_start = paths[0]
                path_x_end = paths[1]
                path_R = paths[2]
                path_graph_last_state = paths[3]
                path_graph_phi_k = paths[4]
                path_IC = paths[5]
                last_state_i = other[1]
                phi_arr_i = other[2]

                mydoc.add_heading('Experiment ' + str(i) + ', ' + str(cycle_num), 1)
                mydoc.add_heading("Initial conditions:", 2)
                for j in range(0, s.k_systems):
                    mydoc.add_paragraph(str(IC_i[j * s.k]) + ', ' + str(IC_i[j * s.k + 1]) + ', ' +
                                        str(IC_i[j * s.k + 2]) + ', ')
                mydoc.add_picture(path_IC, width=docx.shared.Inches(4))

                mydoc.add_picture(path_x_start, width=docx.shared.Inches(6.5))
                mydoc.add_picture(path_x_end, width=docx.shared.Inches(6.5))
                mydoc.add_picture(path_R, width=docx.shared.Inches(4))

                # Рисуем конечное состояние на единичной окружности
                mydoc.add_picture(path_graph_last_state, width=docx.shared.Inches(4))

                # Рисуем график phi(k)
                mydoc.add_picture(path_graph_phi_k, width=docx.shared.Inches(3))

                # Рисуем график изменения суммы фаз относительно первого элемента
                phi_arr_i_T = phi_arr_i.T
                sum_phi = []
                for i in range(len(phi_arr_i[0])):
                    sum_phi.append(sum(phi_arr_i_T[i]))

                plt.figure()
                plt.plot(range(len(sum_phi)) ,sum_phi)
                plt.title('Сумма фаз по модулю 2\u03C0')
                plt.xlabel('k - k-й номер максимума')
                plt.ylabel('sum')
                plt.grid()
                path_graph_sum_phi = s.Graphic_data_path + '_sumPhi_' + str(0) + '.png'
                plt.savefig(path_graph_sum_phi)
                plt.close()
                mydoc.add_picture(path_graph_sum_phi, width=docx.shared.Inches(3))

                mydoc.add_page_break()

                # Сохраняем в док результатов всё что сделали
                mydoc.save(s.path_Doc)

                # Последние элементы массива R1, R2
                R1_arr_last.append(R1_arr_i[-1])
                R2_arr_last.append(R2_arr_i[-1])
                G_inh_arr.append(G_inh_i)

                # Записываем
                mem.write_R12_last(G_inh_arr, R1_arr_last, R2_arr_last, r12_last_path)

        # После каждой череды экспериментов при одной связи строим гистограмму

        plt.figure()
        plt.hist([R1_arr_last, R2_arr_last], color=['blue', 'orange'], 
                         bins=np.linspace(0, 1, 20, endpoint=True), label=['R\u2081', 'R\u2082'])
        plt.legend()
        plt.grid()
        plt.title('Гистограмма при G_inh = ' + str(G_inh_k_exp))
        plt.xlabel('R\u2081, R\u2082')
        # Делаем путь и в отдельную папку кладем графики
        path_hist = './Data/graphics/hist/' + 'hist_' + str(s.k_systems) + 'el_' + 'G' + str(G_inh_k_exp) + '.png'
        plt.savefig(path_hist)
        plt.close()

    
    print('Process end. Final time: ', time.time() - final_time)


#make_protyazhka_R12_dep_G_inh(10, 12, 1)
    
make_protyazhka_R12_dep_G_inh_full(-60, 61, 1, 1000)
make_protyazhka_R12_dep_G_inh_full(60, -61, -1, 1000)
